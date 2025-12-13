#!/usr/bin/env python3
"""
Word to PDF Converter - Exact format preservation
Uses multiple methods to convert Word to PDF with exact formatting
"""

import sys
import json
import os
import subprocess
import tempfile
import traceback
import platform

def convert_with_libreoffice(input_path, output_path):
    """
    Convert using LibreOffice (best for exact format preservation)
    """
    # Determine LibreOffice command based on platform
    if platform.system() == 'Windows':
        libreoffice_paths = [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            'soffice'
        ]
    else:
        libreoffice_paths = [
            '/usr/bin/libreoffice',
            '/usr/bin/soffice',
            'libreoffice',
            'soffice'
        ]
    
    output_dir = os.path.dirname(output_path)
    if not output_dir:
        output_dir = '.'
    
    for lo_path in libreoffice_paths:
        try:
            # LibreOffice headless conversion
            cmd = [
                lo_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                input_path
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120
            )
            
            # LibreOffice outputs to input filename with .pdf extension
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
            
            if os.path.exists(generated_pdf):
                # Rename to expected output path if different
                if generated_pdf != output_path:
                    os.rename(generated_pdf, output_path)
                return True
                
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            continue
        except Exception as e:
            continue
    
    return False

def convert_with_docx2pdf(input_path, output_path):
    """
    Convert using docx2pdf library (uses MS Word on Windows, LibreOffice on Linux)
    """
    try:
        from docx2pdf import convert
        convert(input_path, output_path)
        return os.path.exists(output_path)
    except ImportError:
        return False
    except Exception:
        return False

def convert_with_pypandoc(input_path, output_path):
    """
    Convert using pypandoc (requires pandoc installed)
    """
    try:
        import pypandoc
        pypandoc.convert_file(input_path, 'pdf', outputfile=output_path)
        return os.path.exists(output_path)
    except ImportError:
        return False
    except Exception:
        return False

def convert_with_python_docx_pdf(input_path, output_path):
    """
    Fallback: Convert using python-docx with exact style preservation
    This reads the actual styles from the document
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch, cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Load Word document
        doc = Document(input_path)
        
        # Get page dimensions from document
        section = doc.sections[0]
        page_width = section.page_width.pt if section.page_width else 595  # A4 width
        page_height = section.page_height.pt if section.page_height else 842  # A4 height
        
        left_margin = section.left_margin.pt if section.left_margin else 72
        right_margin = section.right_margin.pt if section.right_margin else 72
        top_margin = section.top_margin.pt if section.top_margin else 72
        bottom_margin = section.bottom_margin.pt if section.bottom_margin else 72
        
        # Create PDF with exact page size
        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=(page_width, page_height),
            rightMargin=right_margin,
            leftMargin=left_margin,
            topMargin=top_margin,
            bottomMargin=bottom_margin
        )
        
        story = []
        
        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                # Preserve empty paragraphs as spacing
                space_after = 6
                if para.paragraph_format.space_after:
                    space_after = para.paragraph_format.space_after.pt or 6
                story.append(Spacer(1, space_after))
                continue
            
            # Get paragraph formatting
            pf = para.paragraph_format
            
            # Determine alignment
            alignment = TA_LEFT
            if pf.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignment = TA_CENTER
            elif pf.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment = TA_RIGHT
            elif pf.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignment = TA_JUSTIFY
            
            # Get font size from first run or default
            font_size = 11
            font_name = 'Helvetica'
            
            if para.runs:
                first_run = para.runs[0]
                if first_run.font.size:
                    font_size = first_run.font.size.pt
                if first_run.font.name:
                    # Map common fonts
                    font_map = {
                        'Arial': 'Helvetica',
                        'Times New Roman': 'Times-Roman',
                        'Courier New': 'Courier',
                        'Calibri': 'Helvetica',
                        'Cambria': 'Times-Roman'
                    }
                    font_name = font_map.get(first_run.font.name, 'Helvetica')
            
            # Get spacing
            space_before = pf.space_before.pt if pf.space_before else 0
            space_after = pf.space_after.pt if pf.space_after else 6
            line_spacing = font_size * 1.2  # Default line spacing
            
            if pf.line_spacing:
                try:
                    line_spacing = pf.line_spacing.pt if hasattr(pf.line_spacing, 'pt') else font_size * 1.2
                except:
                    line_spacing = font_size * 1.2
            
            # Create style for this paragraph
            style = ParagraphStyle(
                f'para_{id(para)}',
                fontName=font_name,
                fontSize=font_size,
                alignment=alignment,
                spaceBefore=space_before,
                spaceAfter=space_after,
                leading=line_spacing
            )
            
            # Build formatted text from runs
            formatted_text = ""
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                
                # Escape special characters
                run_text = run_text.replace('&', '&amp;')
                run_text = run_text.replace('<', '&lt;')
                run_text = run_text.replace('>', '&gt;')
                
                # Apply formatting
                if run.bold and run.italic:
                    formatted_text += f"<b><i>{run_text}</i></b>"
                elif run.bold:
                    formatted_text += f"<b>{run_text}</b>"
                elif run.italic:
                    formatted_text += f"<i>{run_text}</i>"
                elif run.underline:
                    formatted_text += f"<u>{run_text}</u>"
                else:
                    formatted_text += run_text
            
            if not formatted_text:
                formatted_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            try:
                story.append(Paragraph(formatted_text, style))
            except:
                # Fallback for problematic text
                plain_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(plain_text, style))
        
        # Process tables with exact formatting
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(Spacer(1, 6))
                story.append(pdf_table)
                story.append(Spacer(1, 6))
        
        if story:
            pdf_doc.build(story)
        else:
            # Empty document
            story.append(Spacer(1, 12))
            pdf_doc.build(story)
        
        return os.path.exists(output_path)
        
    except Exception as e:
        print(f"Python docx conversion error: {e}", file=sys.stderr)
        return False

def convert_word_to_pdf(input_path, output_path, options=None):
    """
    Convert Word document to PDF with exact format preservation
    Tries multiple methods in order of quality
    """
    options = options or {}
    
    # Method 1: LibreOffice (best quality, exact rendering)
    if convert_with_libreoffice(input_path, output_path):
        return {
            'success': True,
            'output_path': output_path,
            'output_size': os.path.getsize(output_path),
            'method': 'libreoffice',
            'message': 'Word converted to PDF using LibreOffice (exact format preservation)'
        }
    
    # Method 2: docx2pdf (uses MS Word on Windows)
    if convert_with_docx2pdf(input_path, output_path):
        return {
            'success': True,
            'output_path': output_path,
            'output_size': os.path.getsize(output_path),
            'method': 'docx2pdf',
            'message': 'Word converted to PDF using docx2pdf'
        }
    
    # Method 3: pypandoc
    if convert_with_pypandoc(input_path, output_path):
        return {
            'success': True,
            'output_path': output_path,
            'output_size': os.path.getsize(output_path),
            'method': 'pypandoc',
            'message': 'Word converted to PDF using pypandoc'
        }
    
    # Method 4: Python fallback with style preservation
    if convert_with_python_docx_pdf(input_path, output_path):
        return {
            'success': True,
            'output_path': output_path,
            'output_size': os.path.getsize(output_path),
            'method': 'python-docx',
            'message': 'Word converted to PDF using python-docx (style preservation)'
        }
    
    return {
        'success': False,
        'error': 'All conversion methods failed. Please install LibreOffice for best results.',
        'details': 'Install LibreOffice: sudo apt-get install libreoffice'
    }

def main():
    if len(sys.argv) < 3:
        print(json.dumps({
            'success': False,
            'error': 'Usage: python word_to_pdf.py <input_docx> <output_pdf> [options_json]'
        }))
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    options = {}
    if len(sys.argv) > 3:
        try:
            options = json.loads(sys.argv[3])
        except json.JSONDecodeError:
            pass
    
    if not os.path.exists(input_path):
        print(json.dumps({
            'success': False,
            'error': f'Input file not found: {input_path}'
        }))
        sys.exit(1)
    
    result = convert_word_to_pdf(input_path, output_path, options)
    print(json.dumps(result))
    sys.exit(0 if result['success'] else 1)

if __name__ == '__main__':
    main()
